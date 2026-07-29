#!/usr/bin/env python3
"""Extract technological cards and replace the bundled recipe catalogue.

The source directory is treated as read-only. The script creates a generated
TypeScript dataset, updates the two bundled SQLite databases, refreshes the
embedded database module, and writes a compact audit report.
"""

from __future__ import annotations

import argparse
import base64
import json
import re
import shutil
import sqlite3
import subprocess
import unicodedata
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from pypdf import PdfReader


AGE_IDS = (1, 2, 3, 4)
DATASET_VERSION = "2026-07-29-v3"
SKIP_MARKERS = (
    "вихід",
    "выход",
    "маса напівфабрикату",
    "масса полуфабриката",
    "маса відвареної",
    "маса готової",
    "маса страви",
)
SECTION_MARKERS = (
    "технологія приготування",
    "характеристика готової",
    "термін придатності",
    "спосіб реалізації",
    "основні фізичні",
    "харчова (поживна)",
    "карту склав",
    "склав:",
)
INGREDIENT_HEADER_MARKERS = (
    "назва продуктів",
    "назва продукту",
    "найменування сировини",
    "найменування продуктів",
    "вікові групи",
    "маса брутто",
    "маса нетто",
)


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ").replace("\u0007", " ")
    return re.sub(r"\s+", " ", text).strip()


def clean_multiline(value: Any) -> str:
    if value is None:
        return ""
    lines = [clean_text(line) for line in str(value).replace("\xa0", " ").splitlines()]
    return "\n".join(line for line in lines if line)


def normalized_key(value: str) -> str:
    value = unicodedata.normalize("NFC", value).casefold()
    value = re.sub(r"^~\\$", "", value)
    value = re.sub(r"\.(docx|doc|xlsx|pdf)$", "", value)
    value = re.sub(r"^\s*(тк|ттк)\s*", "", value)
    value = re.sub(r"\b(технологічна|технологическая)\s+(карта|картка)\b", " ", value)
    value = re.sub(r"\s*\(\d+\)\s*$", "", value)
    value = re.sub(r"[_‐‑‒–—―]+", " ", value)
    value = re.sub(r"[«»\"'`.,:;№]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    value = value.replace("тюфт", "тефт")
    return value.replace(" ", "")


def as_number(value: Any) -> float | None:
    if value is None or isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return round(float(value), 4)
    text = clean_text(value).replace(",", ".")
    if not text or text in {"-", "—", "–", "⸻"}:
        return None
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return round(float(match.group()), 4) if match else None


def as_total_mass(value: Any) -> float | None:
    text = clean_text(value).replace(",", ".")
    if "/" in text:
        parts = [float(x) for x in re.findall(r"\d+(?:\.\d+)?", text)]
        return round(sum(parts), 4) if parts else None
    return as_number(value)


def unique_text(parts: Iterable[str]) -> str:
    seen: set[str] = set()
    result: list[str] = []
    for part in parts:
        text = clean_multiline(part)
        key = normalized_key(text)
        if text and key not in seen:
            seen.add(key)
            result.append(text)
    return "\n".join(result)


def extract_between(lines: list[str], start_terms: tuple[str, ...], end_terms: tuple[str, ...]) -> str:
    start = None
    collected: list[str] = []
    for index, line in enumerate(lines):
        low = line.casefold()
        if start is None and any(term in low for term in start_terms):
            start = index
            continue
        if start is not None:
            if any(term in low for term in end_terms):
                break
            if line:
                collected.append(line)
    return unique_text(collected)


def title_from_lines(lines: list[str], fallback: str) -> str:
    for index, line in enumerate(lines[:35]):
        low = line.casefold()
        marker = next(
            (item for item in ("технологічна карта", "технологічна картка") if item in low),
            None,
        )
        if marker:
            inline = clean_text(re.split(marker, line, flags=re.IGNORECASE)[-1])
            inline = re.sub(r"^№\s*[\d.]+\s*", "", inline).strip()
            if 2 < len(inline) < 180 and not any(
                term in inline.casefold() for term in ("директор", "затвердж", "рецептур")
            ):
                return inline
            if index + 1 >= len(lines):
                continue
            candidate = clean_text(lines[index + 1])
            if 2 < len(candidate) < 180 and "№ з/п" not in candidate:
                return candidate
    for line in lines[:40]:
        candidate = re.sub(r"^\s*\d+\s*[–—-]\s*", "", clean_text(line))
        letters = [char for char in candidate if char.isalpha()]
        if (
            3 < len(candidate) < 150
            and len(letters) >= 4
            and sum(char.isupper() for char in letters) / len(letters) > 0.8
            and not any(term in candidate.casefold() for term in ("затвердж", "директор", "склав"))
        ):
            return candidate.title()
    for line in lines[:20]:
        candidate = clean_text(line)
        low = candidate.casefold()
        if (
            3 < len(candidate) < 150
            and "затвердж" not in low
            and "рецептур" not in low
            and "збірник" not in low
            and "інститут" not in low
            and "директор" not in low
        ):
            return candidate
    return clean_text(fallback)


def classify_dish(title: str) -> int:
    value = title.casefold()
    rules = [
        (1, ("суп", "борщ", "бульйон", "розсольник", "солянка", "капусняк")),
        (6, ("чай", "какао", "компот", "кисіль", "напій", "сік", "узвар")),
        (5, ("салат", "вінегрет", "ікра овоч", "закуска")),
        (4, ("запікан", "олад", "сирник", "омлет", "пудинг", "биточ", "яйце")),
        (2, ("котлет", "риба", "рибн", "курят", "ялович", "м'яс", "м’яс", "печін", "тефтел", "фрикадел", "гуляш", "шніцел", "рулет")),
        (3, ("каша", "пюре", "картоп", "макарон", "капуста", "рагу", "бобов", "горох", "квасол", "овоч")),
        (7, ("фрукт", "яблук", "банан", "десерт", "мус", "желе", "шарлот", "печив", "булоч", "пиріг")),
    ]
    for category, keywords in rules:
        if any(keyword in value for keyword in keywords):
            return category
    return 2


def product_category(name: str) -> int:
    value = name.casefold()
    rules = [
        (1, ("хліб", "борош", "сухар", "булоч")),
        (2, ("круп", "рис", "макарон", "горох", "квасол", "сочев", "борошно")),
        (3, ("молок", "сметан", "сир", "вершк", "кефір", "йогурт")),
        (4, ("м'яс", "м’яс", "ялович", "кур", "індич", "печін", "ковбас")),
        (5, ("риб", "хек", "минтай", "оселед")),
        (6, ("картоп", "морк", "буряк", "капуст", "цибул", "томат", "огір", "зел", "часник", "овоч", "перець")),
        (7, ("яблу", "банан", "фрукт", "ягод", "родзин", "кураг", "чорнослив", "апельс", "лимон")),
        (8, ("цукор", "мед", "варення", "джем")),
        (9, ("олія", "масло", "маргарин")),
        (10, ("чай", "какао", "сік", "вода")),
        (11, ("яйц",)),
    ]
    for category, keywords in rules:
        if any(keyword in value for keyword in keywords):
            return category
    return 12


def build_metadata(lines: list[str], allergen_codes: list[str]) -> dict[str, str]:
    source = unique_text(
        line for line in lines[:30]
        if any(term in line.casefold() for term in ("рецептур", "збірник", "сборник", "інститут", "видання"))
    )
    allergens = unique_text(
        code for code in allergen_codes
        if code
        and len(code) <= 24
        and "алерген" not in code.casefold()
        and not any(marker in code.casefold() for marker in INGREDIENT_HEADER_MARKERS)
    )
    technology = extract_between(
        lines,
        ("технологія приготування", "технология приготовления"),
        ("характеристика готової", "термін придатності", "основні фізичні", "харчова (поживна)"),
    )
    characteristics = extract_between(
        lines,
        ("характеристика готової",),
        ("основні фізичні", "термін придатності", "харчова (поживна)", "склав:"),
    )
    storage = extract_between(
        lines,
        ("термін придатності",),
        ("спосіб реалізації", "характеристика готової", "харчова (поживна)", "склав:"),
    )
    serving = unique_text(line for line in lines if line.casefold().startswith(("подають", "подається")))
    return {
        "sourceRef": source,
        "allergens": allergens,
        "technology": technology,
        "characteristics": characteristics,
        "storage": storage,
        "serving": serving,
    }


def looks_like_ingredient(name: str) -> bool:
    low = name.casefold()
    if len(name) < 2 or any(marker in low for marker in SKIP_MARKERS):
        return False
    if any(marker in low for marker in SECTION_MARKERS):
        return False
    if any(marker in low for marker in INGREDIENT_HEADER_MARKERS):
        return False
    return not (name.endswith(":") and len(name.split()) < 6)


def ingredient_record(
    name: str,
    gross: list[Any],
    net: list[Any],
    allergen: str = "",
    quality: str = "",
) -> dict[str, Any] | None:
    name = clean_text(name)
    if not looks_like_ingredient(name):
        return None
    gross_numbers = [as_number(value) for value in gross]
    net_numbers = [as_number(value) for value in net]
    if not any(value is not None for value in gross_numbers + net_numbers):
        return None
    for index in range(len(gross_numbers)):
        if gross_numbers[index] is None:
            gross_numbers[index] = net_numbers[index]
        if net_numbers[index] is None:
            net_numbers[index] = gross_numbers[index]
        if isinstance(gross[index], str) and "/" in gross[index] and net_numbers[index] is not None:
            gross_numbers[index] = net_numbers[index]
    return {
        "name": name,
        "allergen": clean_text(allergen),
        "quality": clean_multiline(quality),
        "isAlternative": bool(re.match(r"^\s*(або|или)\b", name, re.IGNORECASE)),
        "gross": gross_numbers,
        "net": net_numbers,
    }


def clone_employee_profile(card: dict[str, Any]) -> None:
    profiles = card["nutrition"]
    if not any(profile["categoryId"] == 4 for profile in profiles):
        source = next((profile for profile in profiles if profile["categoryId"] == 3), None)
        if source:
            profiles.append({**source, "categoryId": 4})
    for ingredient in card["ingredients"]:
        if len(ingredient["gross"]) == 3:
            ingredient["gross"].append(ingredient["gross"][2])
            ingredient["net"].append(ingredient["net"][2])


def parse_wide_rows(
    rows: list[list[Any]],
    merged_word_layout: bool = False,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[str]]:
    ingredients: list[dict[str, Any]] = []
    nutrition: list[dict[str, Any]] = []
    allergens: list[str] = []
    if merged_word_layout:
        gross_columns = (2, 3, 4)
        net_columns = (5, 7, 8)
        yield_columns = (5, 7, 8)
        protein_columns = (9, 11, 12)
        fat_columns = (13, 15, 16)
        carbs_columns = (17, 19, 20)
        calorie_columns = (21, 23, 24)
        quality_column = 25
        minimum_width = 28
    else:
        gross_columns = (2, 3, 4)
        net_columns = (5, 6, 7)
        yield_columns = (5, 6, 7)
        protein_columns = (8, 9, 10)
        fat_columns = (11, 12, 13)
        carbs_columns = (14, 15, 16)
        calorie_columns = (17, 18, 19)
        quality_column = 20
        minimum_width = 21
    for row in rows:
        padded = row + [""] * max(0, minimum_width - len(row))
        name = clean_text(padded[1])
        low = name.casefold()
        if "вихід" in low or "выход" in low:
            for index, category_id in enumerate((1, 2, 3)):
                nutrition.append({
                    "categoryId": category_id,
                    "yieldGr": as_total_mass(padded[yield_columns[index]]) or 0,
                    "protein": as_number(padded[protein_columns[index]]) or 0,
                    "fat": as_number(padded[fat_columns[index]]) or 0,
                    "carbs": as_number(padded[carbs_columns[index]]) or 0,
                    "calories": as_number(padded[calorie_columns[index]]) or 0,
                })
            continue
        allergen = clean_text(padded[0])
        record = ingredient_record(
            name,
            [padded[column] for column in gross_columns],
            [padded[column] for column in net_columns],
            allergen,
            clean_multiline(padded[quality_column]),
        )
        if record:
            ingredients.append(record)
            if allergen and len(allergen) <= 24 and "алерген" not in allergen.casefold():
                allergens.append(allergen)
    return ingredients, nutrition, allergens


def parse_scaled_rows(rows: list[list[Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[str]]:
    output_row = next(
        (
            row for row in rows
            if any(
                "вихідготової" in re.sub(r"\s+", "", clean_text(cell).casefold())
                or "выходготов" in re.sub(r"\s+", "", clean_text(cell).casefold())
                for cell in row
            )
        ),
        None,
    )
    pairs: list[tuple[int, float]] = []
    if output_row:
        for col in range(2, min(len(output_row), 16), 2):
            value = as_total_mass(output_row[col])
            if value:
                pairs.append((col, value))
    pairs = sorted(pairs, key=lambda item: item[1])
    if len(pairs) > 4:
        pairs = pairs[:3] + [pairs[-1]]
    category_pairs = list(zip(AGE_IDS[-len(pairs):] if len(pairs) < 4 else AGE_IDS, pairs))
    ingredients: list[dict[str, Any]] = []
    nutrition = [
        {"categoryId": category_id, "yieldGr": value, "protein": 0, "fat": 0, "carbs": 0, "calories": 0}
        for category_id, (_, value) in category_pairs
    ]
    allergens: list[str] = []
    for row in rows:
        padded = row + [""] * max(0, 21 - len(row))
        name = clean_text(padded[1])
        if clean_text(padded[0]):
            allergens.append(clean_text(padded[0]))
        gross = [padded[col] for _, (col, _) in category_pairs]
        net = [padded[col + 1] for _, (col, _) in category_pairs]
        record = ingredient_record(name, gross, net, clean_text(padded[0]), clean_multiline(padded[20]))
        if record:
            ingredients.append(record)
    while nutrition and len(nutrition) < 4:
        nutrition.insert(0, {**nutrition[0], "categoryId": len(nutrition)})
    return ingredients, nutrition, allergens


def enrich_scaled_nutrition(rows: list[list[Any]], nutrition: list[dict[str, Any]]) -> None:
    for row_index, row in enumerate(rows):
        labels = [clean_text(value).casefold() for value in row]
        if not any("маса порц" in label for label in labels):
            continue
        columns: dict[str, int] = {}
        terms = {
            "yieldGr": ("маса порц",),
            "protein": ("білки", "белки"),
            "fat": ("жири", "жиры"),
            "carbs": ("вуглев", "углев"),
            "calories": ("енергет", "калор"),
        }
        for field, variants in terms.items():
            index = next((i for i, label in enumerate(labels) if any(term in label for term in variants)), None)
            if index is not None:
                columns[field] = index
        if "yieldGr" not in columns:
            continue
        values_by_mass: dict[float, dict[str, float]] = {}
        for data_row in rows[row_index + 1:row_index + 12]:
            mass = as_total_mass(data_row[columns["yieldGr"]] if columns["yieldGr"] < len(data_row) else None)
            if mass is None:
                continue
            values_by_mass[mass] = {
                field: as_number(data_row[column] if column < len(data_row) else None) or 0
                for field, column in columns.items()
                if field != "yieldGr"
            }
        for profile in nutrition:
            values = values_by_mass.get(float(profile["yieldGr"]))
            if values:
                profile.update(values)
        return


def parse_two_age_rows(rows: list[list[Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[str]]:
    ingredients: list[dict[str, Any]] = []
    nutrition: list[dict[str, Any]] = []
    for row in rows:
        padded = row + [""] * max(0, 13 - len(row))
        name = clean_text(padded[0])
        if "вихід" in name.casefold():
            first = {
                "categoryId": 1,
                "yieldGr": as_total_mass(padded[3]) or 0,
                "protein": as_number(padded[5]) or 0,
                "fat": as_number(padded[7]) or 0,
                "carbs": as_number(padded[9]) or 0,
                "calories": as_number(padded[11]) or 0,
            }
            second = {
                "categoryId": 3,
                "yieldGr": as_total_mass(padded[4]) or 0,
                "protein": as_number(padded[6]) or 0,
                "fat": as_number(padded[8]) or 0,
                "carbs": as_number(padded[10]) or 0,
                "calories": as_number(padded[12]) or 0,
            }
            nutrition = [first, {**first, "categoryId": 2}, second, {**second, "categoryId": 4}]
            continue
        record = ingredient_record(
            name,
            [padded[1], padded[1], padded[2], padded[2]],
            [padded[3], padded[3], padded[4], padded[4]],
        )
        if record:
            ingredients.append(record)
    return ingredients, nutrition, []


def spreadsheet_lines(ws: Any, limit: int = 120) -> list[str]:
    lines: list[str] = []
    for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, limit), values_only=True):
        parts = [clean_multiline(value) for value in row if isinstance(value, str) and clean_text(value)]
        if parts:
            lines.append(" ".join(parts))
    return lines


def parse_xlsx_sheet(path: Path, ws: Any, values_ws: Any) -> dict[str, Any]:
    lines = spreadsheet_lines(values_ws)
    rows = [list(row) for row in values_ws.iter_rows(min_row=1, max_row=min(values_ws.max_row, 180), values_only=True)]
    header = " ".join(lines[:20]).casefold()
    if "вікові групи" in header and values_ws.max_column >= 18:
        layout = "wide-age"
        ingredients, nutrition, allergen_codes = parse_wide_rows(rows)
    else:
        layout = "scaled-portions"
        ingredients, nutrition, allergen_codes = parse_scaled_rows(rows)
        enrich_scaled_nutrition(rows, nutrition)
    metadata = build_metadata(lines, allergen_codes)
    card = {
        "title": title_from_lines(lines, ws.title if ws.title.casefold() not in {"лист1", "sheet1"} else path.stem),
        "layout": layout,
        "ingredients": ingredients,
        "nutrition": nutrition,
        **metadata,
    }
    clone_employee_profile(card)
    return card


def parse_xlsx(path: Path) -> list[dict[str, Any]]:
    formula_book = load_workbook(path, read_only=True, data_only=False)
    value_book = load_workbook(path, read_only=True, data_only=True)
    cards: list[dict[str, Any]] = []
    for ws, values_ws in zip(formula_book.worksheets, value_book.worksheets):
        populated = sum(
            1 for row in ws.iter_rows(max_row=min(ws.max_row, 60))
            for cell in row if cell.value is not None
        )
        if populated < 8:
            continue
        cards.append(parse_xlsx_sheet(path, ws, values_ws))
    return cards


def docx_lines(document: Document) -> list[str]:
    lines = [clean_multiline(paragraph.text) for paragraph in document.paragraphs if clean_text(paragraph.text)]
    for table in document.tables:
        for row in table.rows:
            text = unique_text(cell.text for cell in row.cells)
            if text:
                lines.append(text)
    return lines


def iter_docx_blocks(document: DocumentObject) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def parse_docx_segment(path: Path, paragraphs: list[str], tables: list[Table]) -> dict[str, Any]:
    lines = [clean_multiline(text) for text in paragraphs if clean_text(text)]
    for table in tables:
        for row in table.rows:
            text = unique_text(cell.text for cell in row.cells)
            if text:
                lines.append(text)
    wide_rows: list[list[Any]] = []
    scaled_rows: list[list[Any]] = []
    for table in tables:
        rows = [[clean_multiline(cell.text) for cell in row.cells] for row in table.rows]
        if len(table.columns) >= 20:
            wide_rows.extend(rows)
        elif len(table.columns) >= 10:
            scaled_rows.extend(rows)
    joined = " ".join(lines).casefold()
    two_age_rows = [
        [clean_multiline(cell.text) for cell in row.cells]
        for table in tables if len(table.columns) == 13
        for row in table.rows
    ]
    if wide_rows and "вікові групи" in joined:
        layout = "wide-age"
        ingredients, nutrition, allergen_codes = parse_wide_rows(
            wide_rows,
            merged_word_layout=max((len(row) for row in wide_rows), default=0) >= 26,
        )
    elif two_age_rows and "ясла" in joined and "садок" in joined:
        layout = "two-age"
        ingredients, nutrition, allergen_codes = parse_two_age_rows(two_age_rows)
    elif scaled_rows or wide_rows:
        layout = "scaled-portions"
        source_rows = scaled_rows or wide_rows
        ingredients, nutrition, allergen_codes = parse_scaled_rows(source_rows)
        enrich_scaled_nutrition(source_rows, nutrition)
    else:
        layout = "text-only"
        ingredients, nutrition, allergen_codes = [], [], []
    metadata = build_metadata(lines, allergen_codes)
    card = {
        "title": title_from_lines(lines, path.stem),
        "layout": layout,
        "ingredients": ingredients,
        "nutrition": nutrition,
        **metadata,
    }
    clone_employee_profile(card)
    return card


def parse_docx(path: Path) -> list[dict[str, Any]]:
    document = Document(path)
    segments: list[tuple[list[str], list[Table]]] = []
    paragraphs: list[str] = []
    tables: list[Table] = []
    has_recipe_table = False
    for block in iter_docx_blocks(document):
        if isinstance(block, Table):
            block_text = " ".join(clean_text(cell.text) for row in block.rows for cell in row.cells).casefold()
            starts_card = (
                "рецептур" in block_text and ("затвердж" in block_text or len(block.columns) <= 3)
            ) or (
                len(block.columns) <= 3
                and "затвердж" in block_text
                and any(term in block_text for term in ("стор", "видання", "мцфр", "харчуван"))
            )
            if starts_card and has_recipe_table and (paragraphs or tables):
                segments.append((paragraphs, tables))
                paragraphs, tables, has_recipe_table = [], [], False
            tables.append(block)
            if starts_card:
                has_recipe_table = True
        else:
            if clean_text(block.text):
                paragraphs.append(block.text)
    if paragraphs or tables:
        segments.append((paragraphs, tables))
    cards = [parse_docx_segment(path, segment_paragraphs, segment_tables) for segment_paragraphs, segment_tables in segments]
    return cards


def parse_doc(path: Path) -> list[dict[str, Any]]:
    completed = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    raw_text = completed.stdout
    lines = [clean_multiline(line) for line in raw_text.splitlines() if clean_text(line)]
    metadata = build_metadata(lines, [])
    tokens = [clean_text(part) for part in raw_text.split("\x07") if clean_text(part)]
    ingredients: list[dict[str, Any]] = []
    nutrition: list[dict[str, Any]] = []
    output_index = next(
        (index for index, token in enumerate(tokens) if "вихід готової" in token.casefold()),
        None,
    )
    index = 9
    limit = output_index if output_index is not None else len(tokens)
    while index < limit:
        if re.fullmatch(r"\d+", tokens[index]) and index + 1 < limit:
            index += 1
        name = tokens[index]
        values = tokens[index + 1:index + 9]
        if len(values) < 8 or not all(as_number(value) is not None for value in values):
            index += 1
            continue
        record = ingredient_record(
            name,
            [values[2], values[4], values[6], values[6]],
            [values[3], values[5], values[7], values[7]],
        )
        if record:
            ingredients.append(record)
        index += 9
    if output_index is not None:
        output_text = tokens[output_index]
        inline_numbers = [
            as_total_mass(value)
            for value in re.findall(r"\d+(?:[,.]\d+)?", output_text.split("+-")[0])
        ]
        outputs = [value for value in inline_numbers[-4:] if value]
        if len(outputs) < 4:
            outputs = [as_total_mass(token) for token in tokens[output_index + 1:output_index + 6]]
        outputs = [value for value in outputs if value]
        if len(outputs) >= 4:
            outputs = outputs[-3:]
            nutrition = [
                {"categoryId": category_id, "yieldGr": value, "protein": 0, "fat": 0, "carbs": 0, "calories": 0}
                for category_id, value in zip((1, 2, 3), outputs)
            ]
            nutrition.append({**nutrition[-1], "categoryId": 4})
    line_values = [clean_text(line) for line in raw_text.splitlines() if clean_text(line)]
    nutrition_header = next(
        (i for i, line in enumerate(line_values) if line.casefold().startswith("маса порції")),
        None,
    )
    if nutrition_header is not None:
        numeric = [as_number(line) for line in line_values[nutrition_header + 5:]]
        numeric = [value for value in numeric if value is not None]
        rows = [numeric[i:i + 5] for i in range(0, len(numeric), 5) if len(numeric[i:i + 5]) == 5]
        by_mass = {row[0]: row[1:] for row in rows}
        for profile in nutrition:
            values = by_mass.get(profile["yieldGr"])
            if values:
                profile.update({
                    "protein": values[0], "fat": values[1],
                    "carbs": values[2], "calories": values[3],
                })
    return [{
        "title": title_from_lines(lines, path.stem),
        "layout": "legacy-text",
        "ingredients": ingredients,
        "nutrition": nutrition,
        **metadata,
    }]


def parse_pdf(path: Path) -> list[dict[str, Any]]:
    reader = PdfReader(path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    readable = sum(char.isalpha() for char in text) > 200 and (
        sum(char in "іїєґІЇЄҐ" for char in text) > 8
    )
    return [{
        "title": clean_text(path.stem),
        "layout": "pdf-text" if readable else "scanned-pdf",
        "ingredients": [],
        "nutrition": [],
        "sourceRef": "",
        "allergens": "",
        "technology": "",
        "characteristics": "",
        "storage": "",
        "serving": "",
        "needsReview": not readable,
    }]


def parse_file(path: Path) -> list[dict[str, Any]]:
    extension = path.suffix.casefold()
    if extension == ".xlsx":
        cards = parse_xlsx(path)
    elif extension == ".docx":
        cards = parse_docx(path)
    elif extension == ".doc":
        cards = parse_doc(path)
    elif extension == ".pdf":
        cards = parse_pdf(path)
    else:
        raise ValueError(f"Unsupported format: {extension}")
    for index, card in enumerate(cards, 1):
        source_label = path.name if len(cards) == 1 else f"{path.name}#{index}"
        card.update({
            "sourceFile": source_label,
            "sourceFormat": extension.removeprefix(".").upper(),
            "sourceFiles": [source_label],
            "categoryId": classify_dish(card["title"]),
            "needsReview": bool(card.get("needsReview")),
        })
    return cards


def card_score(card: dict[str, Any]) -> int:
    format_score = {"XLSX": 40, "DOCX": 30, "DOC": 10, "PDF": 0}.get(card["sourceFormat"], 0)
    return (
        format_score
        + len(card["ingredients"]) * 3
        + len(card["nutrition"]) * 4
        + bool(card["technology"]) * 10
        + bool(card["allergens"]) * 4
        - bool(card["needsReview"]) * 100
    )


def deduplicate(cards: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for card in cards:
        groups[normalized_key(card["title"])].append(card)
    selected: list[dict[str, Any]] = []
    duplicates: list[dict[str, Any]] = []
    for key, group in groups.items():
        best = max(group, key=card_score)
        best["sourceFiles"] = sorted({card["sourceFile"] for card in group})
        selected.append(best)
        if len(group) > 1:
            duplicates.append({
                "key": key,
                "selected": best["sourceFile"],
                "skipped": sorted(card["sourceFile"] for card in group if card is not best),
            })
    selected.sort(key=lambda card: (card["categoryId"], normalized_key(card["title"])))
    return selected, duplicates


def finalize_cards(cards: list[dict[str, Any]]) -> None:
    for card in cards:
        nutrition = {profile["categoryId"]: profile for profile in card["nutrition"]}
        fallback = nutrition.get(3) or nutrition.get(2) or nutrition.get(1) or {
            "yieldGr": 0, "protein": 0, "fat": 0, "carbs": 0, "calories": 0
        }
        card["defaultNutrition"] = fallback
        for ingredient in card["ingredients"]:
            while len(ingredient["gross"]) < 4:
                ingredient["gross"].append(ingredient["gross"][-1] if ingredient["gross"] else 0)
                ingredient["net"].append(ingredient["net"][-1] if ingredient["net"] else 0)


def update_database(db_path: Path, cards: list[dict[str, Any]], backup_dir: Path) -> None:
    backup_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    shutil.copy2(db_path, backup_dir / f"{db_path.stem}-before-tech-cards-{stamp}.db")
    connection = sqlite3.connect(db_path)
    cursor = connection.cursor()
    cursor.executescript(
        """
        CREATE TABLE IF NOT EXISTS TECH_CARD_NUTRITION (
          ID INTEGER PRIMARY KEY AUTOINCREMENT,
          ID_BLUDA INTEGER NOT NULL,
          ID_KATEGORII_DETEJ INTEGER NOT NULL,
          VYXOD_GR REAL DEFAULT 0,
          BELKI REAL DEFAULT 0,
          ZIRI REAL DEFAULT 0,
          UGLEVODI REAL DEFAULT 0,
          KALORII REAL DEFAULT 0,
          UNIQUE(ID_BLUDA, ID_KATEGORII_DETEJ)
        );
        CREATE TABLE IF NOT EXISTS SADOK_TECH_CARD_IMPORTS (
          DATASET_VERSION TEXT PRIMARY KEY,
          IMPORTED_AT TEXT NOT NULL,
          CARD_COUNT INTEGER NOT NULL,
          SOURCE_COUNT INTEGER NOT NULL
        );
        """
    )
    dish_columns = {row[1] for row in cursor.execute("PRAGMA table_info(KARTOTEKA_BLUD)")}
    for column, declaration in {
        "SOURCE_FILE": "TEXT DEFAULT ''",
        "SOURCE_FORMAT": "TEXT DEFAULT ''",
        "SOURCE_REF": "TEXT DEFAULT ''",
        "ALLERGENS": "TEXT DEFAULT ''",
        "QUALITY_REQUIREMENTS": "TEXT DEFAULT ''",
        "STORAGE_CONDITIONS": "TEXT DEFAULT ''",
        "SERVING_METHOD": "TEXT DEFAULT ''",
        "DISH_CHARACTERISTICS": "TEXT DEFAULT ''",
        "IMPORT_KEY": "TEXT DEFAULT ''",
    }.items():
        if column not in dish_columns:
            cursor.execute(f"ALTER TABLE KARTOTEKA_BLUD ADD COLUMN {column} {declaration}")
    component_columns = {row[1] for row in cursor.execute("PRAGMA table_info(KOMPONENTI_KARTOTEKI)")}
    for column, declaration in {
        "SOURCE_NAME": "TEXT DEFAULT ''",
        "ALLERGENS": "TEXT DEFAULT ''",
        "QUALITY_REQUIREMENTS": "TEXT DEFAULT ''",
        "IS_ALTERNATIVE": "INTEGER DEFAULT 0",
    }.items():
        if column not in component_columns:
            cursor.execute(f"ALTER TABLE KOMPONENTI_KARTOTEKI ADD COLUMN {column} {declaration}")

    product_rows = cursor.execute("SELECT ID, NAME FROM PRODUKTS").fetchall()
    product_map = {normalized_key(name): product_id for product_id, name in product_rows}
    old_menu = cursor.execute("SELECT ID, NAME_BLUDA FROM MENU").fetchall()
    cursor.execute("DELETE FROM TECH_CARD_NUTRITION")
    cursor.execute("DELETE FROM KOMPONENTI_KARTOTEKI")
    cursor.execute("DELETE FROM KARTOTEKA_BLUD")
    cursor.execute("DELETE FROM sqlite_sequence WHERE name IN ('KARTOTEKA_BLUD','KOMPONENTI_KARTOTEKI','TECH_CARD_NUTRITION')")

    dish_ids: dict[str, int] = {}
    for order, card in enumerate(cards, 1):
        quality = unique_text(item["quality"] for item in card["ingredients"] if item["quality"])
        default = card["defaultNutrition"]
        cursor.execute(
            """
            INSERT INTO KARTOTEKA_BLUD
              (NAME, NOTES, ID_GRUPPI_BLUD, VYXOD, BELKI, ZIRI, UGLEVODI, KALORII,
               PORRDOK_SLEDOVANIR_BLUD, SOURCE_FILE, SOURCE_FORMAT, SOURCE_REF,
               ALLERGENS, QUALITY_REQUIREMENTS, STORAGE_CONDITIONS, SERVING_METHOD,
               DISH_CHARACTERISTICS, IMPORT_KEY)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                card["title"], card["technology"], card["categoryId"],
                default["yieldGr"], default["protein"], default["fat"],
                default["carbs"], default["calories"], order,
                card["sourceFile"], card["sourceFormat"], card["sourceRef"],
                card["allergens"], quality, card["storage"], card["serving"],
                card["characteristics"], normalized_key(card["title"]),
            ),
        )
        dish_id = cursor.lastrowid
        dish_ids[normalized_key(card["title"])] = dish_id
        for profile in card["nutrition"]:
            cursor.execute(
                """
                INSERT OR REPLACE INTO TECH_CARD_NUTRITION
                  (ID_BLUDA, ID_KATEGORII_DETEJ, VYXOD_GR, BELKI, ZIRI, UGLEVODI, KALORII)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    dish_id, profile["categoryId"], profile["yieldGr"], profile["protein"],
                    profile["fat"], profile["carbs"], profile["calories"],
                ),
            )
        for line, ingredient in enumerate(card["ingredients"], 1):
            product_key = normalized_key(ingredient["name"])
            product_id = product_map.get(product_key)
            if product_id is None:
                cursor.execute(
                    """
                    INSERT INTO PRODUKTS
                      (NAME, ID_GRUPPI_PRODUKTOV, EDINICA_IZMERENIA, NOMER_PP)
                    VALUES (?, ?, 'кг', ?)
                    """,
                    (ingredient["name"], product_category(ingredient["name"]), len(product_map) + 1),
                )
                product_id = cursor.lastrowid
                product_map[product_key] = product_id
            for category_index, category_id in enumerate(AGE_IDS):
                cursor.execute(
                    """
                    INSERT INTO KOMPONENTI_KARTOTEKI
                      (ID_BLUDA, ID_PRODUKTA, ID_KATEGORII_DETEJ, GROSSO_GR, NETTO_GR,
                       NOMER_ID_LINII_V_TABLICE, SOURCE_NAME, ALLERGENS,
                       QUALITY_REQUIREMENTS, IS_ALTERNATIVE)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        dish_id, product_id, category_id,
                        ingredient["gross"][category_index] or 0,
                        ingredient["net"][category_index] or 0,
                        line, ingredient["name"], ingredient["allergen"],
                        ingredient["quality"], int(ingredient["isAlternative"]),
                    ),
                )
    for menu_id, name in old_menu:
        dish_id = dish_ids.get(normalized_key(name or ""))
        if dish_id:
            cursor.execute("UPDATE MENU SET ID_BLUDA = ? WHERE ID = ?", (dish_id, menu_id))
    cursor.execute("DELETE FROM SADOK_TECH_CARD_IMPORTS")
    cursor.execute(
        "INSERT INTO SADOK_TECH_CARD_IMPORTS VALUES (?, ?, ?, ?)",
        (DATASET_VERSION, datetime.now(timezone.utc).isoformat(), len(cards), sum(len(card["sourceFiles"]) for card in cards)),
    )
    connection.commit()
    status = cursor.execute("PRAGMA integrity_check").fetchone()[0]
    if status != "ok":
        raise RuntimeError(f"SQLite integrity check failed for {db_path}: {status}")
    connection.close()


def write_ts_dataset(path: Path, cards: list[dict[str, Any]]) -> None:
    payload = json.dumps(cards, ensure_ascii=False, separators=(",", ":"))
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        "// Generated by scripts/import_tech_cards.py. Do not edit manually.\n"
        f"export const TECH_CARD_DATASET_VERSION = {json.dumps(DATASET_VERSION)};\n"
        f"export const IMPORTED_TECH_CARDS: any[] = {payload};\n",
        encoding="utf-8",
    )


def write_embedded_database(db_path: Path, target: Path) -> None:
    encoded = base64.b64encode(db_path.read_bytes()).decode("ascii")
    target.write_text(
        "// System generated embedded SQLite database fallback\n"
        f'const DB_BASE64 = "{encoded}";\n\n'
        "export function getEmbeddedDbBytes(): Uint8Array {\n"
        "  const binaryString = atob(DB_BASE64);\n"
        "  const bytes = new Uint8Array(binaryString.length);\n"
        "  for (let i = 0; i < binaryString.length; i++) {\n"
        "    bytes[i] = binaryString.charCodeAt(i);\n"
        "  }\n"
        "  return bytes;\n"
        "}\n",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=Path("картки"))
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    source = args.source
    paths = sorted(
        path for path in source.iterdir()
        if path.is_file()
        and path.suffix.casefold() in {".xlsx", ".docx", ".doc", ".pdf"}
        and not path.name.startswith("~$")
    )
    parsed: list[dict[str, Any]] = []
    errors: list[dict[str, str]] = []
    for path in paths:
        try:
            parsed.extend(parse_file(path))
        except Exception as error:
            errors.append({"file": path.name, "error": f"{type(error).__name__}: {error}"})
    cards, duplicates = deduplicate(parsed)
    finalize_cards(cards)
    needs_review = [
        {
            "file": card["sourceFile"],
            "title": card["title"],
            "reason": "Скан PDF без надёжного текстового слоя"
            if card["layout"] == "scanned-pdf"
            else "Не удалось извлечь состав или возрастные нормы",
        }
        for card in cards
        if card["needsReview"] or not card["ingredients"] or not card["nutrition"]
    ]
    report = {
        "sourceFiles": len(paths),
        "parsedFiles": len(parsed),
        "importedCards": len(cards),
        "skippedDuplicates": sum(len(item["skipped"]) for item in duplicates),
        "ingredients": sum(len(card["ingredients"]) for card in cards),
        "nutritionProfiles": sum(len(card["nutrition"]) for card in cards),
        "errors": errors,
        "needsReview": needs_review,
        "duplicates": duplicates,
    }
    Path("tech_cards_import_report.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps({key: value for key, value in report.items() if key not in {"duplicates"}}, ensure_ascii=False, indent=2))
    if args.dry_run:
        return
    write_ts_dataset(Path("src/data/importedTechCards.ts"), cards)
    backup_dir = Path("backups/tech-cards")
    update_database(Path("medsestra.db"), cards, backup_dir)
    shutil.copy2(Path("medsestra.db"), Path("public/medsestra.db"))
    write_embedded_database(Path("medsestra.db"), Path("src/services/db_data.ts"))


if __name__ == "__main__":
    main()
